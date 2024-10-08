import os
from flask import Flask, render_template, request
from docx import Document  
from docx.shared import Pt
import io
import fitz  
import joblib
import pandas as pd

# -------------------------------------------------------------------
# MetaDocAI - Dokumen Pemeriksaan dan Validasi Otomatis
# Versi: 1.3
# Hak Cipta © 2024 Fapzarz
# -------------------------------------------------------------------

app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

# Memuat model dan kolom fitur ML jika tersedia
model = None
model_columns = None
if os.path.exists('document_classifier.joblib') and os.path.exists('model_columns.joblib'):
    try:
        model = joblib.load('document_classifier.joblib')
        model_columns = joblib.load('model_columns.joblib')
    except Exception as e:
        print("Gagal memuat model ML:", e)

def check_docx(doc: Document) -> dict:
    report = []
    success = True

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name != 'Times New Roman':
                report.append(f'Font tidak sesuai di paragraf: "{para.text[:30]}..."')
                success = False
                break
            if run.font.size and run.font.size.pt != 12:
                report.append(f'Ukuran font tidak sesuai di paragraf: "{para.text[:30]}..."')
                success = False
                break

        if para.paragraph_format.line_spacing is not None:
            spacing = para.paragraph_format.line_spacing
            if spacing != 1.5:
                report.append(f'Spasi tidak sesuai di paragraf: "{para.text[:30]}..."')
                success = False

    sections = doc.sections
    if sections:
        section = sections[0]

        left_margin_cm = section.left_margin.inches * 2.54
        right_margin_cm = section.right_margin.inches * 2.54
        top_margin_cm = section.top_margin.inches * 2.54
        bottom_margin_cm = section.bottom_margin.inches * 2.54

        expected_margins_cm = {
            'left': 4.0,    # 4 cm
            'right': 3.0,   # 3 cm
            'top': 3.0,     # 3 cm
            'bottom': 3.0   # 3 cm
        }

        tolerance = 0.1  # cm

        if abs(left_margin_cm - expected_margins_cm['left']) > tolerance:
            report.append(f'Margin kiri tidak sesuai: {left_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["left"]} cm)')
            success = False
        if abs(right_margin_cm - expected_margins_cm['right']) > tolerance:
            report.append(f'Margin kanan tidak sesuai: {right_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["right"]} cm)')
            success = False
        if abs(top_margin_cm - expected_margins_cm['top']) > tolerance:
            report.append(f'Margin atas tidak sesuai: {top_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["top"]} cm)')
            success = False
        if abs(bottom_margin_cm - expected_margins_cm['bottom']) > tolerance:
            report.append(f'Margin bawah tidak sesuai: {bottom_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["bottom"]} cm)')
            success = False
    else:
        report.append('Tidak dapat memeriksa margin.')
        success = False

    result = {"success": success, "messages": report}

    # Tambahkan prediksi ML jika model tersedia
    if model and model_columns:
        try:
            features = extract_features_docx(doc)
            features_df = prepare_features(features)
            prediction = model.predict(features_df)[0]
            prediction_prob = model.predict_proba(features_df)[0][1]  # Probabilitas kelas '1' (Correct)

            ml_message = "Dokumen sesuai dengan kriteria (Prediksi ML: Correct)." if prediction == 1 else "Dokumen tidak sesuai dengan kriteria (Prediksi ML: Incorrect)."
            result["messages"].append(ml_message)
            result["confidence"] = f"{prediction_prob * 100:.2f}%"
            result["success"] = result["success"] and (prediction == 1)
        except Exception as e:
            result["messages"].append("Gagal melakukan prediksi ML pada dokumen.")
            result["success"] = False

    return result

def check_pdf(file_stream: io.BytesIO) -> dict:
    report = []
    success = True
    try:
        doc = fitz.open(stream=file_stream.read(), filetype="pdf")
    except Exception as e:
        report.append('Gagal membaca dokumen PDF.')
        return {"success": False, "messages": report}

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block['type'] != 0:  
                continue
            for line in block.get('lines', []):
                for span in line.get('spans', []):
                    font = span.get('font', '').lower()
                    size = span.get('size', 0)
                    text = span.get('text', '').strip()

                    if 'times new roman' not in font:
                        report.append(f'Font tidak sesuai di halaman {page_num +1}: "{text[:30]}..."')
                        success = False
                    if abs(size - 12) > 0.5:  
                        report.append(f'Ukuran font tidak sesuai di halaman {page_num +1}: "{text[:30]}..."')
                        success = False

    report.append('Pemeriksaan margin pada PDF tidak dilakukan secara mendalam.')

    result = {"success": success and not report, "messages": report}

    # Tambahkan prediksi ML jika model tersedia
    if model and model_columns:
        try:
            features = extract_features_pdf(file_stream)
            features_df = prepare_features(features)
            prediction = model.predict(features_df)[0]
            prediction_prob = model.predict_proba(features_df)[0][1]

            ml_message = "Dokumen sesuai dengan kriteria (Prediksi ML: Correct)." if prediction == 1 else "Dokumen tidak sesuai dengan kriteria (Prediksi ML: Incorrect)."
            result["messages"].append(ml_message)
            result["confidence"] = f"{prediction_prob * 100:.2f}%"
            result["success"] = result["success"] and (prediction == 1)
        except Exception as e:
            result["messages"].append("Gagal melakukan prediksi ML pada dokumen PDF.")
            result["success"] = False

    return result

def extract_features_docx(doc: Document) -> dict:
    features = {}
    fonts = {}
    sizes = {}
    line_spacings = set()
    margins = {}

    for para in doc.paragraphs:
        for run in para.runs:
            font = run.font.name or "Unknown"
            size = run.font.size.pt if run.font.size else 12
            fonts[font] = fonts.get(font, 0) + 1
            sizes[size] = sizes.get(size, 0) + 1
        if para.paragraph_format.line_spacing:
            line_spacings.add(para.paragraph_format.line_spacing)

    # Margins
    if doc.sections:
        section = doc.sections[0]
        margins['left_margin_cm'] = section.left_margin.inches * 2.54
        margins['right_margin_cm'] = section.right_margin.inches * 2.54
        margins['top_margin_cm'] = section.top_margin.inches * 2.54
        margins['bottom_margin_cm'] = section.bottom_margin.inches * 2.54
    else:
        margins = {'left_margin_cm': 0, 'right_margin_cm': 0, 'top_margin_cm': 0, 'bottom_margin_cm': 0}

    # Fitur tambahan
    features['unique_fonts'] = len(fonts)
    features['most_common_font'] = max(fonts, key=fonts.get) if fonts else "Unknown"
    features['font_size_variance'] = pd.Series(list(sizes.keys())).std() if sizes else 0
    features['unique_line_spacings'] = len(line_spacings)
    # Margins
    for key, value in margins.items():
        features[key] = value

    return features

def extract_features_pdf(file_stream: io.BytesIO) -> dict:
    features = {}
    try:
        doc = fitz.open(stream=file_stream.read(), filetype="pdf")
    except Exception as e:
        print("Gagal membuka PDF:", e)
        return features

    fonts = {}
    sizes = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block['type'] != 0:
                continue
            for line in block.get('lines', []):
                for span in line.get('spans', []):
                    font = span.get('font', 'Unknown')
                    size = span.get('size', 12)
                    fonts[font] = fonts.get(font, 0) + 1
                    sizes.append(size)

    features['unique_fonts'] = len(fonts)
    features['most_common_font'] = max(fonts, key=fonts.get) if fonts else "Unknown"
    features['font_size_variance'] = pd.Series(sizes).std() if sizes else 0
    # Margins pada PDF sulit dieksktrak, jadi kita set ke 0
    features['left_margin_cm'] = 0
    features['right_margin_cm'] = 0
    features['top_margin_cm'] = 0
    features['bottom_margin_cm'] = 0

    return features

def prepare_features(features: dict) -> pd.DataFrame:
    df = pd.DataFrame([features])
    df = pd.get_dummies(df, columns=['most_common_font'], drop_first=True)

    # Menambahkan kolom yang mungkin hilang sesuai model
    for col in model_columns:
        if col not in df.columns:
            df[col] = 0

    df = df[model_columns]
    df.fillna(0, inplace=True)

    return df

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        uploaded_files = request.files.getlist('files')
        if not uploaded_files or uploaded_files == [None]:
            return render_template('index.html', reports=[{"success": False, "filename": "Tidak ada file yang diunggah.", "messages": ["Tidak ada file yang diunggah."]}])

        reports = []
        for uploaded_file in uploaded_files:
            if uploaded_file.filename == '':
                continue  # Lewati file kosong

            file_name = uploaded_file.filename.lower()
            if not (file_name.endswith('.docx') or file_name.endswith('.pdf')):
                reports.append({
                    "success": False,
                    "filename": uploaded_file.filename,
                    "messages": ["Silakan unggah file .docx atau .pdf saja."]
                })
                continue

            file_bytes = uploaded_file.read()
            file_stream = io.BytesIO(file_bytes)

            if file_name.endswith('.docx'):
                try:
                    doc = Document(file_stream)
                    report = check_docx(doc)
                except Exception as e:
                    report = {"success": False, "messages": ["Gagal membaca dokumen .docx. Pastikan file dalam format yang benar."]}
            elif file_name.endswith('.pdf'):
                report = check_pdf(file_stream)

            report["filename"] = uploaded_file.filename
            reports.append(report)

        return render_template('index.html', reports=reports)

    return render_template('index.html')

@app.errorhandler(413)
def request_entity_too_large(error):
    return render_template('index.html', reports=[{"success": False, "filename": "File terlalu besar.", "messages": ["File terlalu besar. Maksimal 50MB."]}]), 413

if __name__ == '__main__':
    print("Isi folder 'templates':", os.listdir('templates'))
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 81)))
